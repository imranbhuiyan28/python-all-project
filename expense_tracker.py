from docx import Document
from prettytable import PrettyTable


def main():
    expenses = []
    try:
        balance = float(input("Enter the Budget: "))

        # if not balance.replace('.', '', 1).isdigit():
        #     print("Invalid balance input. Please enter a valid number.")
        #     return
        # balance = float(balance)
    except ValueError:
        print("Enter valid number")
        return

    print("Welcome to expense tracker app")
    print("Your total balance is: ", balance)
    while True:
        print("Menu:")
        print("1. Add Expense")
        print("2. View Expenses")
        print("3. View Remaining Budget")
        print("4. Exit")
        print("5. Save to Docx")
        user_input = input("Enter the option for access: ")

        if user_input == '4':
            print("Thank you for using our system, Good bye!")
            break
        elif user_input == '5':
            save_to_word(expenses, balance)
        elif user_input == '1':
            balance, expenses = add_expense(balance, expenses)
        elif user_input == '2':
            view_expenses(expenses, balance)
        elif user_input == '3':
            remaining_balance(balance)
        else:
            print("Invalid input, try again")


def add_expense(balance, expenses):
    while True:
        description = input('Enter expense description (or type exit to stop): ')

        if description.lower() == 'exit':
            print("Exiting..... Your remaining balance is: ", balance)
            save_to_word(expenses, balance)
            break

        try:
            expense_amount = float(input('Enter the expense amount: '))

            if expense_amount <= 0:
                print("Expense cannot be 0 or negative. Try again with a valid input.")
            elif expense_amount > balance:
                print("Insufficient balance.")
            else:
                balance -= expense_amount
                expenses.append({"Description": description, "Amount": expense_amount})
                print("Your expense has been added. Type 2 to view expenses.")

        except ValueError:
            print("Invalid input! Please enter a valid number.")

    return balance, expenses


def view_expenses(expenses, balance):
    if not expenses:
        print("No expenses recorded yet.")
    else:
        # Create PrettyTable to display the expenses in a tabular format
        table = PrettyTable()
        table.field_names = ["Description", "Amount"]

        for expense in expenses:
            table.add_row([expense["Description"], expense["Amount"]])

        table.add_row(["Remaining Balance: ", balance])

        print("\nExpense List:")
        print(table)


def remaining_balance(balance):
    print("Your remaining balance: ", balance)


def save_to_word(expenses, balance):
    doc = Document()
    doc.add_heading("Expense Tracker Report")
    doc.add_paragraph(f"Remaining Budget: {balance}\n")

    if not expenses:
        doc.add_paragraph("No expenses recorded.")
    else:
        doc.add_heading("Total expenses:", level=2)

        # Create a table in the Word document using PrettyTable formatted data
        table = PrettyTable()
        table.field_names = ["Description", "Amount"]

        for expense in expenses:
            table.add_row([expense["Description"], expense["Amount"]])

        # Add the formatted table to the Word document
        doc.add_paragraph(str(table))

    # Save the file
    file_name = "Expense_tracker.docx"
    doc.save(file_name)
    print(f"Your file has been saved as {file_name}")


# Run the main program
main()
